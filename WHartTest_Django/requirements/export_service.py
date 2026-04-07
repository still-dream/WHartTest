"""
评审报告导出服务
支持导出为Excel、Word、PDF格式
"""
import io
import os
from datetime import datetime
from typing import Optional
from xml.sax.saxutils import escape
from django.http import HttpResponse
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from .models import ReviewReport, ReviewIssue, ModuleReviewResult


# issue_type 显示名称映射
ISSUE_TYPE_DISPLAY_MAP = {
    'specification': '规范性',
    'clarity': '清晰度',
    'completeness': '完整性',
    'consistency': '一致性',
    'feasibility': '可行性',
    'logic': '逻辑性',
    'testability': '可测性',
}

# 优先级显示名称映射
PRIORITY_DISPLAY_MAP = {
    'high': '高',
    'medium': '中',
    'low': '低',
}

ANALYSIS_TYPE_META = [
    ('completeness_analysis', '完整性分析'),
    ('consistency_analysis', '一致性分析'),
    ('testability_analysis', '可测性分析'),
    ('feasibility_analysis', '可行性分析'),
    ('clarity_analysis', '清晰度分析'),
    ('logic_analysis', '逻辑分析'),
]


def _pdf_safe_text(value) -> str:
    """转义 PDF Paragraph 文本，避免特殊字符导致渲染异常。"""
    return escape(str(value or '')).replace('\n', '<br/>')


def _bullet_lines(items) -> str:
    """将列表转为带项目符号的 HTML 文本。"""
    valid_items = [str(item).strip() for item in (items or []) if str(item).strip()]
    return '<br/>'.join([f'• {_pdf_safe_text(item)}' for item in valid_items])


def _get_issues_from_specialized_analyses(report: ReviewReport) -> list[dict]:
    """
    从 specialized_analyses JSON 中提取所有问题列表。
    
    返回统一格式的问题字典列表，每个包含：
    - title: 问题标题（从description截取或用category）
    - issue_type: 问题类型（source字段）
    - issue_type_display: 类型中文名
    - priority: 优先级
    - priority_display: 优先级中文名
    - description: 详细描述
    - suggestion: 改进建议
    - location: 位置
    - category: 分类
    """
    sa = report.specialized_analyses or {}
    if not isinstance(sa, dict) or not sa:
        return []
    
    all_issues = []
    # 按 source 排序：completeness, consistency, clarity, feasibility, logic, testability
    order_map = {'completeness_analysis': 0, 'consistency_analysis': 1, 'clarity_analysis': 2,
                 'feasibility_analysis': 3, 'logic_analysis': 4, 'testability_analysis': 5}
    
    sorted_keys = sorted(sa.keys(), key=lambda k: order_map.get(k, 99))
    
    for analysis_key in sorted_keys:
        analysis_data = sa.get(analysis_key)
        if not isinstance(analysis_data, dict):
            continue
        
        issues_list = analysis_data.get('issues', [])
        if not isinstance(issues_list, list):
            continue
        
        for iss in issues_list:
            if not isinstance(iss, dict):
                continue
            
            source = iss.get('source', '')
            desc = iss.get('description', '') or ''
            
            # 从描述生成标题：取前40个字符
            title = iss.get('title', '') or ''
            if not title:
                if len(desc) > 40:
                    title = desc[:40] + '...'
                else:
                    title = desc
            
            # 映射显示值
            issue_type_display = ISSUE_TYPE_DISPLAY_MAP.get(source, source)
            priority_val = iss.get('priority', 'medium')
            priority_display = PRIORITY_DISPLAY_MAP.get(priority_val, priority_val)
            category = iss.get('category', '')
            
            # 组合位置信息
            location_parts = []
            if iss.get('location'):
                location_parts.append(iss['location'])
            if iss.get('module_name'):
                location_parts.append(iss['module_name'])
            location = ' | '.join(location_parts) if location_parts else (iss.get('section') or '')
            
            all_issues.append({
                'title': title.strip(),
                'issue_type': source,
                'issue_type_display': issue_type_display,
                'priority': priority_val,
                'priority_display': priority_display,
                'description': desc,
                'suggestion': iss.get('suggestion', '') or '',
                'location': location,
                'category': category,
                'analysis_key': analysis_key.replace('_analysis', ''),
            })
    
    return all_issues


class ReportExportService:
    """评审报告导出服务"""
    
    # 中文字体路径（需要确保系统中安装了中文字体）
    # 对于Docker环境，通常在 /usr/share/fonts/ 下
    CHINESE_FONT_PATHS = [
        '/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc',  # 文泉驿
        '/usr/share/fonts/truetype/wqy/wqy-microhei.ttc',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',  # Noto Sans CJK
        'C:\\Windows\\Fonts\\simhei.ttf',  # Windows黑体
        'C:\\Windows\\Fonts\\msyh.ttc',  # Windows微软雅黑
        '/System/Library/Fonts/PingFang.ttc',  # macOS
    ]
    
    @classmethod
    def find_chinese_font(cls):
        """查找可用的中文字体"""
        for font_path in cls.CHINESE_FONT_PATHS:
            if os.path.exists(font_path):
                return font_path
        return None
    
    @classmethod
    def export_to_excel(cls, report: ReviewReport) -> HttpResponse:
        """导出报告为Excel格式"""
        workbook = Workbook()
        
        # 删除默认sheet
        default_sheet = workbook.active
        workbook.remove(default_sheet)
        
        # 创建概览sheet
        cls._create_overview_sheet(workbook, report)
        
        # 创建问题列表sheet
        cls._create_issues_sheet(workbook, report)
        
        # 创建模块评审结果sheet
        cls._create_modules_sheet(workbook, report)
        
        # 创建专项分析sheet
        cls._create_specialized_sheet(workbook, report)
        
        # 保存到字节流
        output = io.BytesIO()
        workbook.save(output)
        output.seek(0)
        
        # 创建响应
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        filename = f'评审报告_{report.document.title}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    
    @classmethod
    def export_to_word(cls, report: ReviewReport) -> HttpResponse:
        """导出报告为Word格式"""
        document = Document()
        
        # 标题
        title = document.add_heading(f'{report.document.title} - 评审报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 基本信息
        document.add_heading('一、评审基本信息', level=1)
        info_table = document.add_table(rows=6, cols=2)
        info_table.style = 'Light Grid Accent 1'
        
        info_data = [
            ('文档标题', report.document.title),
            ('评审时间', report.review_date.strftime('%Y-%m-%d %H:%M')),
            ('评审人', report.reviewer),
            ('总体评价', report.get_overall_rating_display() or '未评级'),
            ('完整度评分', f'{report.completion_score}分'),
            ('问题总数', f'{report.total_issues}个'),
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)
        
        # 评审总结
        document.add_heading('二、评审总结', level=1)
        document.add_paragraph(report.summary or '无')
        
        if report.recommendations:
            document.add_heading('改进建议', level=2)
            document.add_paragraph(report.recommendations)
        
        # 问题统计
        document.add_heading('三、问题统计', level=1)
        stats_table = document.add_table(rows=4, cols=2)
        stats_table.style = 'Light Grid Accent 1'
        
        stats_data = [
            ('高优先级问题', f'{report.high_priority_issues}个'),
            ('中优先级问题', f'{report.medium_priority_issues}个'),
            ('低优先级问题', f'{report.low_priority_issues}个'),
            ('问题总数', f'{report.total_issues}个'),
        ]
        
        for i, (label, value) in enumerate(stats_data):
            stats_table.rows[i].cells[0].text = label
            stats_table.rows[i].cells[1].text = value
        
        # 详细问题列表
        issues = report.issues.all()
        if issues:
            document.add_heading('四、详细问题列表', level=1)
            
            for idx, issue in enumerate(issues, 1):
                # 问题标题
                issue_title = document.add_heading(
                    f'{idx}. {issue.title}',
                    level=2
                )
                
                # 问题信息表格
                issue_table = document.add_table(rows=5, cols=2)
                issue_table.style = 'Light Grid Accent 1'
                
                issue_data = [
                    ('问题类型', issue.get_issue_type_display()),
                    ('优先级', issue.get_priority_display()),
                    ('位置', issue.location or '未指定'),
                    ('状态', '已解决' if issue.is_resolved else '未解决'),
                    ('问题描述', issue.description),
                ]
                
                for i, (label, value) in enumerate(issue_data):
                    issue_table.rows[i].cells[0].text = label
                    issue_table.rows[i].cells[1].text = value
                
                if issue.suggestion:
                    document.add_paragraph(f'改进建议：{issue.suggestion}')
                
                document.add_paragraph()  # 空行
        
        # 模块评审结果
        module_results = report.module_results.all()
        if module_results:
            document.add_heading('五、模块评审结果', level=1)
            
            for result in module_results:
                module_title = document.add_heading(result.module.title, level=2)
                
                result_table = document.add_table(rows=4, cols=2)
                result_table.style = 'Light Grid Accent 1'
                
                result_data = [
                    ('模块评价', result.get_module_rating_display() or '未评级'),
                    ('问题数量', f'{result.issues_count}个'),
                    ('严重程度', f'{result.severity_score}分'),
                ]
                
                for i, (label, value) in enumerate(result_data):
                    result_table.rows[i].cells[0].text = label
                    result_table.rows[i].cells[1].text = value
                
                if result.strengths:
                    document.add_paragraph(f'优点：{result.strengths}')
                if result.weaknesses:
                    document.add_paragraph(f'不足：{result.weaknesses}')
                if result.recommendations:
                    document.add_paragraph(f'改进建议：{result.recommendations}')
                
                document.add_paragraph()
        
        # 评分详情
        document.add_heading('六、评分详情', level=1)
        scores_table = document.add_table(rows=7, cols=2)
        scores_table.style = 'Light Grid Accent 1'
        
        scores_data = [
            ('完整性评分', f'{report.completeness_score}分'),
            ('一致性评分', f'{report.consistency_score}分'),
            ('可测性评分', f'{report.testability_score}分'),
            ('可行性评分', f'{report.feasibility_score}分'),
            ('清晰度评分', f'{report.clarity_score}分'),
            ('逻辑分析评分', f'{report.logic_score}分'),
            ('完整度评分', f'{report.completion_score}分'),
        ]
        
        for i, (label, value) in enumerate(scores_data):
            scores_table.rows[i].cells[0].text = label
            scores_table.rows[i].cells[1].text = value
        
        # 保存到字节流
        output = io.BytesIO()
        document.save(output)
        output.seek(0)
        
        # 创建响应
        response = HttpResponse(
            output.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        filename = f'评审报告_{report.document.title}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
    
    @classmethod
    def export_to_pdf(cls, report: ReviewReport) -> HttpResponse:
        """导出报告为PDF格式（按专项分析页面样式输出）"""
        output = io.BytesIO()

        doc = SimpleDocTemplate(
            output,
            pagesize=A4,
            rightMargin=48,
            leftMargin=48,
            topMargin=48,
            bottomMargin=48
        )

        font_path = cls.find_chinese_font()
        if font_path:
            try:
                pdfmetrics.registerFont(TTFont('Chinese', font_path))
                chinese_font = 'Chinese'
            except Exception:
                chinese_font = 'Helvetica'
        else:
            chinese_font = 'Helvetica'

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'PdfTitle',
            parent=styles['Heading1'],
            fontName=chinese_font,
            fontSize=20,
            alignment=1,
            textColor=colors.HexColor('#1D2129'),
            spaceAfter=18,
        )
        section_style = ParagraphStyle(
            'PdfSection',
            parent=styles['Heading2'],
            fontName=chinese_font,
            fontSize=15,
            textColor=colors.HexColor('#1D2129'),
            spaceBefore=12,
            spaceAfter=8,
        )
        card_title_style = ParagraphStyle(
            'PdfCardTitle',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=11,
            leading=16,
            textColor=colors.HexColor('#1D2129'),
        )
        card_body_style = ParagraphStyle(
            'PdfCardBody',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=10,
            leading=16,
            textColor=colors.HexColor('#4E5969'),
        )
        score_style = ParagraphStyle(
            'PdfScore',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=24,
            leading=28,
            alignment=1,
            textColor=colors.white,
        )
        score_label_style = ParagraphStyle(
            'PdfScoreLabel',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=10,
            leading=12,
            alignment=1,
            textColor=colors.white,
        )
        issue_meta_style = ParagraphStyle(
            'PdfIssueMeta',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=9,
            leading=14,
            textColor=colors.HexColor('#86909C'),
        )
        issue_title_style = ParagraphStyle(
            'PdfIssueTitle',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=11,
            leading=16,
            textColor=colors.HexColor('#1D2129'),
        )
        issue_desc_style = ParagraphStyle(
            'PdfIssueDesc',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=10,
            leading=16,
            textColor=colors.HexColor('#4E5969'),
        )
        issue_suggestion_style = ParagraphStyle(
            'PdfIssueSuggestion',
            parent=styles['Normal'],
            fontName=chinese_font,
            fontSize=10,
            leading=16,
            textColor=colors.HexColor('#165DFF'),
        )

        story = []
        story.append(Paragraph(f'{_pdf_safe_text(report.document.title)} - 评审报告', title_style))

        info_data = [
            [Paragraph('<b>文档标题</b>', card_title_style), Paragraph(_pdf_safe_text(report.document.title), card_body_style)],
            [Paragraph('<b>评审时间</b>', card_title_style), Paragraph(_pdf_safe_text(report.review_date.strftime('%Y-%m-%d %H:%M')), card_body_style)],
            [Paragraph('<b>评审人</b>', card_title_style), Paragraph(_pdf_safe_text(report.reviewer), card_body_style)],
            [Paragraph('<b>总体评价</b>', card_title_style), Paragraph(_pdf_safe_text(report.get_overall_rating_display() or '未评级'), card_body_style)],
            [Paragraph('<b>完整度评分</b>', card_title_style), Paragraph(_pdf_safe_text(f'{report.completion_score}分'), card_body_style)],
            [Paragraph('<b>问题总数</b>', card_title_style), Paragraph(_pdf_safe_text(f'{report.total_issues}个'), card_body_style)],
        ]
        info_table = Table(info_data, colWidths=[1.4 * inch, 4.4 * inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F2F3F5')),
            ('BACKGROUND', (1, 0), (1, -1), colors.white),
            ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#E5E6EB')),
            ('INNERGRID', (0, 0), (-1, -1), 0.8, colors.HexColor('#E5E6EB')),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(info_table)
        story.append(Spacer(1, 12))

        specialized_analyses = report.specialized_analyses if isinstance(report.specialized_analyses, dict) else {}
        analysis_entries = [
            (analysis_key, analysis_title, specialized_analyses.get(analysis_key))
            for analysis_key, analysis_title in ANALYSIS_TYPE_META
            if isinstance(specialized_analyses.get(analysis_key), dict)
        ]

        if not analysis_entries:
            story.append(Paragraph('暂无专项分析数据', card_body_style))
        else:
            for index, (analysis_key, analysis_title, analysis_data) in enumerate(analysis_entries, 1):
                overall_score = analysis_data.get('overall_score', 0)
                summary = analysis_data.get('summary') or '暂无分析总结'
                strengths = analysis_data.get('strengths') if isinstance(analysis_data.get('strengths'), list) else []
                recommendations = analysis_data.get('recommendations') if isinstance(analysis_data.get('recommendations'), list) else []
                issues = analysis_data.get('issues') if isinstance(analysis_data.get('issues'), list) else []

                story.append(Paragraph(f'{index}. {_pdf_safe_text(analysis_title)}', section_style))

                score_paragraph = Paragraph(
                    f'<para align="center"><font size="22"><b>{overall_score}</b></font><br/><font size="10">分</font></para>',
                    score_style
                )
                summary_paragraph = Paragraph(
                    f'<b>分析总结</b><br/>{_pdf_safe_text(summary)}',
                    card_body_style
                )
                summary_table = Table([[score_paragraph, summary_paragraph]], colWidths=[1.15 * inch, 4.75 * inch])
                summary_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (0, 0), colors.HexColor('#1890FF')),
                    ('BACKGROUND', (1, 0), (1, 0), colors.HexColor('#F7F8FA')),
                    ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#E5E6EB')),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 12),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                    ('TOPPADDING', (0, 0), (-1, -1), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ]))
                story.append(summary_table)
                story.append(Spacer(1, 8))

                if strengths:
                    strengths_table = Table([[Paragraph(f'<b>优势</b><br/>{_bullet_lines(strengths)}', card_body_style)]], colWidths=[5.9 * inch])
                    strengths_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F6FFED')),
                        ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#B7EB8F')),
                        ('LEFTPADDING', (0, 0), (-1, -1), 12),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                        ('TOPPADDING', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                    ]))
                    story.append(strengths_table)
                    story.append(Spacer(1, 8))

                if recommendations:
                    recommendations_table = Table([[Paragraph(f'<b>改进建议</b><br/>{_bullet_lines(recommendations)}', card_body_style)]], colWidths=[5.9 * inch])
                    recommendations_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#E8F4FF')),
                        ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#91CAFF')),
                        ('LEFTPADDING', (0, 0), (-1, -1), 12),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                        ('TOPPADDING', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                    ]))
                    story.append(recommendations_table)
                    story.append(Spacer(1, 8))

                story.append(Paragraph(f'发现的问题 ({len(issues)}个)', card_title_style))
                story.append(Spacer(1, 4))

                if issues:
                    for issue in issues:
                        priority_raw = str(issue.get('severity') or issue.get('priority') or 'medium').lower()
                        priority_display = PRIORITY_DISPLAY_MAP.get(priority_raw, priority_raw)
                        category = issue.get('category') or '未分类'
                        location = issue.get('location') or issue.get('section') or ''
                        title = issue.get('title') or issue.get('description') or '未命名问题'
                        description = issue.get('description') or ''
                        suggestion = issue.get('suggestion') or ''

                        meta_parts = [f'[{priority_display}]', _pdf_safe_text(category)]
                        if location:
                            meta_parts.append(f'位置：{_pdf_safe_text(location)}')
                        meta_text = '　'.join(meta_parts)

                        issue_rows = [
                            [Paragraph(meta_text, issue_meta_style)],
                            [Paragraph(f'<b>{_pdf_safe_text(title)}</b>', issue_title_style)],
                        ]
                        if description and description != title:
                            issue_rows.append([Paragraph(_pdf_safe_text(description), issue_desc_style)])
                        if suggestion:
                            issue_rows.append([Paragraph(f'<b>建议：</b>{_pdf_safe_text(suggestion)}', issue_suggestion_style)])

                        issue_table = Table(issue_rows, colWidths=[5.9 * inch])
                        issue_style = [
                            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFF7E8')),
                            ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#E5E6EB')),
                            ('LINEBEFORE', (0, 0), (0, -1), 3, colors.HexColor('#1890FF')),
                            ('LEFTPADDING', (0, 0), (-1, -1), 10),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 10),
                            ('TOPPADDING', (0, 0), (-1, -1), 8),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                        ]
                        if suggestion:
                            suggestion_row_index = len(issue_rows) - 1
                            issue_style.append(('BACKGROUND', (0, suggestion_row_index), (-1, suggestion_row_index), colors.HexColor('#E8F4FF')))
                        issue_table.setStyle(TableStyle(issue_style))
                        story.append(issue_table)
                        story.append(Spacer(1, 8))
                else:
                    empty_table = Table([[Paragraph('该维度暂无发现问题', card_body_style)]], colWidths=[5.9 * inch])
                    empty_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F7F8FA')),
                        ('BOX', (0, 0), (-1, -1), 0.8, colors.HexColor('#E5E6EB')),
                        ('LEFTPADDING', (0, 0), (-1, -1), 12),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 12),
                        ('TOPPADDING', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                    ]))
                    story.append(empty_table)

                story.append(Spacer(1, 12))

        doc.build(story)
        output.seek(0)

        response = HttpResponse(
            output.getvalue(),
            content_type='application/pdf'
        )
        filename = f'评审报告_{report.document.title}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'

        return response
    
    @classmethod
    def _create_overview_sheet(cls, workbook: Workbook, report: ReviewReport):
        """创建概览sheet"""
        sheet = workbook.create_sheet('评审概览', 0)
        
        # 标题样式
        title_font = Font(name='微软雅黑', size=16, bold=True, color='FFFFFF')
        title_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        title_alignment = Alignment(horizontal='center', vertical='center')
        
        # 内容样式
        header_font = Font(name='微软雅黑', size=11, bold=True)
        header_fill = PatternFill(start_color='D9E2F3', end_color='D9E2F3', fill_type='solid')
        content_font = Font(name='微软雅黑', size=11)
        content_alignment = Alignment(horizontal='left', vertical='center', wrap_text=True)
        
        # 边框样式
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 标题
        sheet.merge_cells('A1:F1')
        cell = sheet['A1']
        cell.value = f'{report.document.title} - 评审报告'
        cell.font = title_font
        cell.fill = title_fill
        cell.alignment = title_alignment
        sheet.row_dimensions[1].height = 30
        
        # 基本信息
        info_data = [
            ['评审时间', report.review_date.strftime('%Y-%m-%d %H:%M'), '', '评审人', report.reviewer, ''],
            ['总体评价', report.get_overall_rating_display() or '未评级', '', '完整度评分', f'{report.completion_score}分', ''],
            ['问题总数', f'{report.total_issues}个', '', '高优先级', f'{report.high_priority_issues}个', ''],
            ['中优先级', f'{report.medium_priority_issues}个', '', '低优先级', f'{report.low_priority_issues}个', ''],
        ]
        
        start_row = 3
        for i, row_data in enumerate(info_data):
            row_num = start_row + i
            for j, value in enumerate(row_data):
                cell = sheet.cell(row=row_num, column=j+1, value=value)
                cell.font = header_font if j % 2 == 0 else content_font
                cell.fill = header_fill if j % 2 == 0 else PatternFill()
                cell.alignment = content_alignment
                cell.border = thin_border
        
        # 调整列宽
        sheet.column_dimensions['A'].width = 15
        sheet.column_dimensions['B'].width = 20
        sheet.column_dimensions['C'].width = 5
        sheet.column_dimensions['D'].width = 15
        sheet.column_dimensions['E'].width = 20
        sheet.column_dimensions['F'].width = 5
        
        # 评审总结
        start_row = 8
        sheet.merge_cells(f'A{start_row}:F{start_row}')
        cell = sheet[f'A{start_row}']
        cell.value = '评审总结'
        cell.font = Font(name='微软雅黑', size=14, bold=True, color='FFFFFF')
        cell.fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        cell.alignment = Alignment(horizontal='left', vertical='center')
        sheet.row_dimensions[start_row].height = 25
        
        sheet.merge_cells(f'A{start_row+1}:F{start_row+1}')
        cell = sheet[f'A{start_row+1}']
        cell.value = report.summary or '无'
        cell.font = content_font
        cell.alignment = content_alignment
        cell.border = thin_border
        sheet.row_dimensions[start_row+1].height = 60
        
        # 改进建议
        if report.recommendations:
            start_row = 11
            sheet.merge_cells(f'A{start_row}:F{start_row}')
            cell = sheet[f'A{start_row}']
            cell.value = '改进建议'
            cell.font = Font(name='微软雅黑', size=14, bold=True, color='FFFFFF')
            cell.fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
            cell.alignment = Alignment(horizontal='left', vertical='center')
            
            sheet.merge_cells(f'A{start_row+1}:F{start_row+1}')
            cell = sheet[f'A{start_row+1}']
            cell.value = report.recommendations
            cell.font = content_font
            cell.alignment = content_alignment
            cell.border = thin_border
            sheet.row_dimensions[start_row+1].height = 60
    
    @classmethod
    def _create_issues_sheet(cls, workbook: Workbook, report: ReviewReport):
        """创建问题列表sheet"""
        sheet = workbook.create_sheet('问题列表', 1)
        
        # 样式定义
        header_font = Font(name='微软雅黑', size=11, bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        
        content_font = Font(name='微软雅黑', size=10)
        content_alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
        
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 表头
        headers = ['序号', '问题标题', '问题类型', '优先级', '位置', '状态', '问题描述', '改进建议']
        for i, header in enumerate(headers):
            cell = sheet.cell(row=1, column=i+1, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        
        sheet.row_dimensions[1].height = 25
        
        # 数据行
        issues = report.issues.all().order_by('priority', '-created_at')
        for idx, issue in enumerate(issues, 1):
            row_num = idx + 1
            
            row_data = [
                idx,
                issue.title,
                issue.get_issue_type_display(),
                issue.get_priority_display(),
                issue.location or '未指定',
                '已解决' if issue.is_resolved else '未解决',
                issue.description,
                issue.suggestion or '',
            ]
            
            for j, value in enumerate(row_data):
                cell = sheet.cell(row=row_num, column=j+1, value=value)
                cell.font = content_font
                cell.alignment = content_alignment
                cell.border = thin_border
                
                # 根据优先级设置颜色
                if j == 3:  # 优先级列
                    if issue.priority == 'high':
                        cell.fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
                    elif issue.priority == 'medium':
                        cell.fill = PatternFill(start_color='FFEB9C', end_color='FFEB9C', fill_type='solid')
                    else:
                        cell.fill = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
            
            sheet.row_dimensions[row_num].height = 40
        
        # 调整列宽
        column_widths = [6, 25, 12, 10, 20, 10, 40, 40]
        for i, width in enumerate(column_widths):
            sheet.column_dimensions[get_column_letter(i+1)].width = width
    
    @classmethod
    def _create_modules_sheet(cls, workbook: Workbook, report: ReviewReport):
        """创建模块评审结果sheet"""
        sheet = workbook.create_sheet('模块评审结果', 2)
        
        # 样式定义
        header_font = Font(name='微软雅黑', size=11, bold=True, color='FFFFFF')
        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        
        content_font = Font(name='微软雅黑', size=10)
        content_alignment = Alignment(horizontal='left', vertical='top', wrap_text=True)
        
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 表头
        headers = ['序号', '模块名称', '评价', '问题数', '严重度', '优点', '不足', '改进建议']
        for i, header in enumerate(headers):
            cell = sheet.cell(row=1, column=i+1, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
            cell.border = thin_border
        
        sheet.row_dimensions[1].height = 25
        
        # 数据行
        module_results = report.module_results.all().order_by('module__order')
        for idx, result in enumerate(module_results, 1):
            row_num = idx + 1
            
            row_data = [
                idx,
                result.module.title,
                result.get_module_rating_display() or '未评级',
                result.issues_count,
                f'{result.severity_score}分',
                result.strengths or '',
                result.weaknesses or '',
                result.recommendations or '',
            ]
            
            for j, value in enumerate(row_data):
                cell = sheet.cell(row=row_num, column=j+1, value=value)
                cell.font = content_font
                cell.alignment = content_alignment
                cell.border = thin_border
            
            sheet.row_dimensions[row_num].height = 60
        
        # 调整列宽
        column_widths = [6, 20, 12, 10, 10, 30, 30, 30]
        for i, width in enumerate(column_widths):
            sheet.column_dimensions[get_column_letter(i+1)].width = width
    
    @classmethod
    def _create_specialized_sheet(cls, workbook: Workbook, report: ReviewReport):
        """创建专项分析sheet"""
        sheet = workbook.create_sheet('专项分析评分', 3)
        
        # 样式定义
        title_font = Font(name='微软雅黑', size=14, bold=True, color='FFFFFF')
        title_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
        
        header_font = Font(name='微软雅黑', size=11, bold=True)
        header_fill = PatternFill(start_color='D9E2F3', end_color='D9E2F3', fill_type='solid')
        
        content_font = Font(name='微软雅黑', size=11)
        
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 标题
        sheet.merge_cells('A1:D1')
        cell = sheet['A1']
        cell.value = '专项分析评分'
        cell.font = title_font
        cell.fill = title_fill
        cell.alignment = Alignment(horizontal='center', vertical='center')
        sheet.row_dimensions[1].height = 30
        
        # 评分数据
        scores = [
            ('完整性评分', report.completeness_score),
            ('一致性评分', report.consistency_score),
            ('可测性评分', report.testability_score),
            ('可行性评分', report.feasibility_score),
            ('清晰度评分', report.clarity_score),
            ('逻辑分析评分', report.logic_score),
            ('完整度评分', report.completion_score),
        ]
        
        # 表头
        headers = ['评分项', '分数', '满分', '百分比']
        for i, header in enumerate(headers):
            cell = sheet.cell(row=2, column=i+1, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = thin_border
        
        # 数据行
        for idx, (name, score) in enumerate(scores, 3):
            sheet.cell(row=idx, column=1, value=name).font = content_font
            sheet.cell(row=idx, column=2, value=score).font = content_font
            sheet.cell(row=idx, column=3, value=100).font = content_font
            percentage = f'{score}%'
            sheet.cell(row=idx, column=4, value=percentage).font = content_font
            
            for j in range(1, 5):
                cell = sheet.cell(row=idx, column=j)
                cell.border = thin_border
                cell.alignment = Alignment(horizontal='center', vertical='center')
        
        # 调整列宽
        sheet.column_dimensions['A'].width = 20
        sheet.column_dimensions['B'].width = 12
        sheet.column_dimensions['C'].width = 12
        sheet.column_dimensions['D'].width = 12
